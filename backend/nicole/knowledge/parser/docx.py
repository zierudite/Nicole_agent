"""DOCX 解析器。

使用 python-docx 解析 Word 文档，提取段落、表格、样式。
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from .base import BaseParser, ParserResult

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Word 文档解析器 (python-docx)。"""

    def __init__(
        self,
        use_paddlex: bool = False,
        use_rapidocr: bool = False,
    ):
        self.use_paddlex = use_paddlex
        self.use_rapidocr = use_rapidocr

    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    async def parse(self, file_path: str) -> Optional[ParserResult]:
        """解析 .docx 文件。"""
        if not self.validate_file(file_path):
            return None

        path = Path(file_path)
        metadata = {"filename": path.name, "size": path.stat().st_size}

        try:
            import docx
            from docx.document import Document

            doc = docx.Document(file_path)
            paragraphs = []
            tables_text = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    style = para.style.name if para.style else ""
                    prefix = f"[{style}] " if style and style != "Normal" else ""
                    paragraphs.append(f"{prefix}{para.text}")

            # 提取表格
            for i, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))
                tables_text.append(f"--- Table {i + 1} ---\n" + "\n".join(table_rows))

            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)

            return ParserResult(
                text=full_text,
                file_type="docx",
                metadata={
                    **metadata,
                    "paragraphs": len(paragraphs),
                    "tables": len(tables_text),
                    "engine": "python-docx",
                },
            )

        except ImportError:
            logger.error("python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"DOCX parse failed: {e}")
            return None
